"""Create ATS-friendly Word resumes from editable plain-text content."""

from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)


def parse_resume(text: str) -> list[tuple[str, str]]:
    """Parse the intentionally simple editable CV syntax."""
    blocks: list[tuple[str, str]] = []
    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            blocks.append(("space", ""))
        elif line.startswith("### "):
            blocks.append(("heading3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("heading2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
        elif line.startswith(('- ', '* ', '• ')):
            blocks.append(("bullet", line[2:].strip()))
        else:
            blocks.append(("body", line))
    return blocks


def _set_cell_free_font(run, size=11, color=None, bold=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color or RGBColor(0, 0, 0)
    run.bold = bold


def _add_bottom_border(paragraph, color="D9E2F3", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def build_docx(resume_text: str) -> bytes:
    """Build a polished single-column Word resume in memory."""
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 15, BLUE, 12, 5),
        ("Heading 2", 12.5, BLUE, 10, 4),
        ("Heading 3", 11.5, DARK_BLUE, 7, 3),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    first_title = True
    blocks = parse_resume(resume_text)
    for index, (kind, content) in enumerate(blocks):
        if kind == "space":
            continue
        if kind == "title":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(content)
            _set_cell_free_font(run, size=23, color=DARK_BLUE, bold=True)
            if first_title:
                _add_bottom_border(paragraph)
                first_title = False
        elif kind == "heading2":
            document.add_paragraph(content, style="Heading 1")
        elif kind == "heading3":
            document.add_paragraph(content, style="Heading 2")
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.first_line_indent = Inches(-0.16)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.keep_with_next = (
                index + 1 < len(blocks) and blocks[index + 1][0] == "bullet"
            )
            run = paragraph.add_run(content)
            _set_cell_free_font(run, size=10.5)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(content)
            _set_cell_free_font(run, size=10.5, color=MUTED if "@" in content else None)

    candidate_name = next(
        (content for kind, content in blocks if kind == "title"), "Candidate"
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{candidate_name} | CV")
    _set_cell_free_font(run, size=8.5, color=MUTED)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(resume_text: str) -> bytes:
    """Build a polished single-column PDF resume in memory."""
    output = BytesIO()
    candidate_name = next(
        (content for kind, content in parse_resume(resume_text) if kind == "title"),
        "Candidate",
    )
    document = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title=f"{candidate_name} - CV",
        author=candidate_name,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CVBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5,
        leading=12.2, spaceAfter=4, textColor=colors.HexColor("#111827"), alignment=TA_LEFT,
    )
    title = ParagraphStyle(
        "CVTitle", parent=body, fontName="Helvetica-Bold", fontSize=23,
        leading=27, textColor=colors.HexColor("#1F4D78"), spaceAfter=7,
        borderColor=colors.HexColor("#D9E2F3"), borderWidth=0,
        borderPadding=(0, 0, 5, 0),
    )
    heading2 = ParagraphStyle(
        "CVH2", parent=body, fontName="Helvetica-Bold", fontSize=15,
        leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=5,
        keepWithNext=True,
    )
    heading3 = ParagraphStyle(
        "CVH3", parent=body, fontName="Helvetica-Bold", fontSize=12.5,
        leading=15, textColor=colors.HexColor("#1F4D78"), spaceBefore=7, spaceAfter=4,
        keepWithNext=True,
    )
    bullet_style = ParagraphStyle(
        "CVBullet", parent=body, leftIndent=16, firstLineIndent=0, spaceAfter=3,
    )

    story = []
    blocks = parse_resume(resume_text)

    def pdf_bullet(content):
        return ListFlowable(
            [ListItem(Paragraph(escape(content), bullet_style), leftIndent=0)],
            bulletType="bullet", start="circle", leftIndent=15, bulletFontSize=7,
            spaceAfter=1,
        )

    index = 0
    while index < len(blocks):
        kind, content = blocks[index]
        safe = escape(content)
        if kind == "space":
            continue
        if kind == "title":
            story.append(Paragraph(safe, title))
        elif kind == "heading2":
            story.append(Paragraph(safe, heading2))
        elif kind == "heading3":
            role_block = [Paragraph(safe, heading3)]
            next_index = index + 1
            while next_index < len(blocks) and blocks[next_index][0] == "bullet":
                role_block.append(pdf_bullet(blocks[next_index][1]))
                next_index += 1
            story.append(KeepTogether(role_block))
            index = next_index - 1
        elif kind == "bullet":
            story.append(pdf_bullet(content))
        else:
            story.append(Paragraph(safe, body))
        index += 1

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#667085"))
        canvas.drawCentredString(
            LETTER[0] / 2, 0.35 * inch, f"{candidate_name} | CV | {doc.page}"
        )
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()
