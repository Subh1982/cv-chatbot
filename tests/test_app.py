import json
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from pypdf import PdfReader

from app import (
    APP_TITLE,
    DEFAULT_MODEL,
    FALLBACK_MESSAGE,
    answer_question,
    assess_job_fit,
    extract_cv_text,
    get_api_key,
    safe_export_name,
    suggest_cv_improvements,
    validate_job_url,
    validate_cv_edits,
)
from resume_export import build_docx, build_pdf


def fake_client(payload):
    response = SimpleNamespace(text=json.dumps(payload) if payload is not None else "bad json")
    models = SimpleNamespace(generate_content=lambda **_: response)
    return SimpleNamespace(models=models)


@patch("app.genai.Client")
def test_returns_grounded_answer(client):
    client.return_value = fake_client({"supported": True, "answer": "19+ years."})
    assert answer_question("Experience?", "19 years", "test") == "19+ years."


@patch("app.genai.Client")
def test_unsupported_uses_exact_fallback(client):
    client.return_value = fake_client({"supported": False, "answer": ""})
    assert answer_question("Favourite colour?", "CV", "test") == FALLBACK_MESSAGE


@patch("app.genai.Client")
def test_invalid_output_uses_exact_fallback(client):
    client.return_value = fake_client(None)
    assert answer_question("Question", "CV", "test") == FALLBACK_MESSAGE


@patch("app.genai.Client")
def test_empty_supported_answer_uses_exact_fallback(client):
    client.return_value = fake_client({"supported": True, "answer": "  "})
    assert answer_question("Question", "CV", "test") == FALLBACK_MESSAGE


@patch.dict("app.os.environ", {}, clear=True)
@patch("app.st.secrets")
def test_missing_secrets_file_is_optional(secrets):
    secrets.get.side_effect = FileNotFoundError
    assert get_api_key() == ""


@patch.dict("app.os.environ", {"GEMINI_API_KEY": " env-key "}, clear=True)
def test_environment_key_does_not_read_secrets():
    assert get_api_key() == "env-key"


def test_default_model_is_available_to_new_users():
    assert DEFAULT_MODEL == "gemini-3.5-flash-lite"


def test_app_is_generic():
    assert APP_TITLE == "CV-JD Compatibility Checker"


def test_fallback_is_friendly_and_actionable():
    assert "uploaded CV does not clearly provide an answer" in FALLBACK_MESSAGE


def test_job_url_must_be_complete_public_web_url():
    assert validate_job_url(" https://example.com/job ") == "https://example.com/job"
    for invalid in ("", "example.com/job", "file:///tmp/job", "javascript:alert(1)"):
        try:
            validate_job_url(invalid)
            assert False, f"Expected {invalid!r} to be rejected"
        except ValueError:
            pass


@patch("app.genai.Client")
def test_job_assessment_uses_url_context_and_clamps_score(client):
    payload = {
        "accessible": True,
        "candidate_name": "Alex Candidate",
        "job_title": "Product Manager",
        "company": "Example Co",
        "score": 105,
        "justification": "Strong product leadership alignment.",
        "strengths": ["Product strategy"],
        "gaps": ["Industry experience is not stated"],
    }
    response = SimpleNamespace(text=json.dumps(payload))
    calls = []
    client.return_value = SimpleNamespace(
        models=SimpleNamespace(generate_content=lambda **kwargs: calls.append(kwargs) or response)
    )

    result = assess_job_fit("https://example.com/job", "CV text", "test")

    assert result["score"] == 100
    assert calls[0]["config"].tools[0].url_context is not None


@patch("app.genai.Client")
def test_inaccessible_job_page_has_clear_error(client):
    payload = {
        "accessible": False,
        "candidate_name": "",
        "job_title": "",
        "company": "",
        "score": 0,
        "justification": "",
        "strengths": [],
        "gaps": [],
    }
    client.return_value = fake_client(payload)
    try:
        assess_job_fit("https://example.com/job", "CV", "test")
        assert False, "Expected inaccessible URL to fail"
    except ValueError as exc:
        assert "could not be read" in str(exc)


@patch("app.genai.Client")
def test_pasted_job_description_does_not_use_url_context(client):
    payload = {
        "accessible": True,
        "candidate_name": "Alex Candidate",
        "job_title": "Senior Product Manager",
        "company": "Example Co",
        "score": 85,
        "justification": "Strong match.",
        "strengths": ["Leadership"],
        "gaps": ["One requirement is not evidenced"],
    }
    response = SimpleNamespace(text=json.dumps(payload))
    calls = []
    client.return_value = SimpleNamespace(
        models=SimpleNamespace(generate_content=lambda **kwargs: calls.append(kwargs) or response)
    )
    description = "Senior Product Manager role. " + ("Responsibilities and requirements. " * 5)

    result = assess_job_fit(
        "https://www.linkedin.com/jobs/view/4453530466/",
        "CV text",
        "test",
        job_description=description,
    )

    assert result["score"] == 85
    assert calls[0]["config"].tools is None
    assert description.strip() in calls[0]["contents"]


def test_short_pasted_job_description_is_rejected():
    try:
        assess_job_fit("", "CV", "test", job_description="Too short")
        assert False, "Expected short description to fail"
    except ValueError as exc:
        assert "full job description" in str(exc)


@patch("app.genai.Client")
def test_tailored_cv_generation_uses_only_pasted_job_context(client):
    payload = {
        "suggestions": ["Prioritise relevant product leadership outcomes."],
        "tailored_cv": "# Alex Candidate\n## Experience\n- Led product strategy.",
    }
    response = SimpleNamespace(text=json.dumps(payload))
    calls = []
    client.return_value = SimpleNamespace(
        models=SimpleNamespace(generate_content=lambda **kwargs: calls.append(kwargs) or response)
    )
    job = "Senior Product Manager. " + ("Responsibilities and requirements. " * 5)

    result = suggest_cv_improvements("", job, "Original CV", "test")

    assert result["tailored_cv"].startswith("# Alex")
    assert calls[0]["config"].tools is None
    assert "Never invent" in calls[0]["contents"]


@patch("app.genai.Client")
def test_validation_blocks_unsupported_cv_claims(client):
    client.return_value = fake_client(
        {"valid": False, "issues": ["The edited CV adds an unsupported award."]}
    )
    result = validate_cv_edits("Edited CV " * 60, "Original CV", "test")
    assert result["valid"] is False
    assert "unsupported award" in result["issues"][0]


def test_short_edited_cv_is_blocked_without_api_call():
    assert validate_cv_edits("Too short", "Original", "test")["valid"] is False


def test_docx_and_pdf_exports_are_valid_documents():
    text = (
        "# Alex Candidate\n"
        "Email: alex@example.com\n"
        "## Profile\nProduct leader with digital experience.\n"
        "## Experience\n### Senior Product Manager\n"
        "- Led product strategy and measurable delivery outcomes.\n"
        "- Managed cross-functional product teams.\n"
        "## Education\n- Master of Business Administration"
    )
    docx_bytes = build_docx(text)
    pdf_bytes = build_pdf(text)

    assert docx_bytes.startswith(b"PK")
    doc = Document(BytesIO(docx_bytes))
    assert "Alex Candidate" in "\n".join(p.text for p in doc.paragraphs)
    assert pdf_bytes.startswith(b"%PDF")
    pdf = PdfReader(BytesIO(pdf_bytes))
    assert "Alex Candidate" in (pdf.pages[0].extract_text() or "")


def test_docx_cv_upload_is_readable():
    source = Document()
    source.add_heading("Alex Candidate", level=1)
    source.add_paragraph("Product manager with digital delivery experience.")
    stream = BytesIO()
    source.save(stream)
    stream.seek(0)
    stream.name = "alex-cv.docx"

    extracted = extract_cv_text(stream)

    assert "Alex Candidate" in extracted
    assert "digital delivery experience" in extracted


def test_export_filename_is_safe():
    assert safe_export_name("Alex Candidate", "Senior Product Manager / AI") == (
        "Alex_Candidate_CV_Senior_Product_Manager_AI.docx"
    )
